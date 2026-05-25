#!/usr/bin/env python3
"""
generate_textseries.py — TextSeries 新製品ジェネレーター

使い方:
  python shared/gtm/scripts/generate_textseries.py \
    --slug rifotext \
    --industry "リフォーム会社" \
    --docs "工事見積書,施主向け提案書,工事完了報告書" \
    --emoji "🔨" \
    --rank 17 \
    [--product-name RifoText] \
    [--no-stripe] \
    [--dry-run]

生成されるファイル:
  C:/Users/ryuuM/fudotext/saas-dev/projects/{slug}/app.py
  C:/Users/ryuuM/fudotext/saas-dev/projects/{slug}/gen.py
  C:/Users/ryuuM/fudotext/saas-dev/projects/{slug}/db.py
  C:/Users/ryuuM/fudotext/saas-dev/projects/{slug}/requirements.txt
  C:/Users/ryuuM/ai-holdings/shared/gtm/config/{slug}.json
  C:/Users/ryuuM/ai-holdings/docs/{slug}.html
  C:/Users/ryuuM/ai-holdings/.github/workflows/{slug}-*.yml  (5本)
  C:/Users/ryuuM/ai-holdings/shared/gtm/outreach/templates/{slug}_sequence_*.txt
  C:/Users/ryuuM/ai-holdings/saas-dev/projects/{slug}/outreach/  (ディレクトリ)
  C:/Users/ryuuM/ai-holdings/saas-dev/projects/{slug}/feedback_reports/  (ディレクトリ)
"""

import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import argparse
import json
import os
import subprocess
import urllib.request
import urllib.error
from pathlib import Path

AI_HOLDINGS = Path(__file__).resolve().parent.parent.parent.parent
FUDOTEXT = Path("C:/Users/ryuuM/fudotext")


# ── Gemini helper ──────────────────────────────────────────────────────────────

def _gemini(prompt: str, api_key: str, max_tokens: int = 6000) -> str:
    payload = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": max_tokens, "temperature": 0.4},
    }).encode()
    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-2.5-flash:generateContent?key={api_key}"
    )
    req = urllib.request.Request(
        url, data=payload, headers={"Content-Type": "application/json"}, method="POST"
    )
    with urllib.request.urlopen(req, timeout=90) as r:
        data = json.loads(r.read())
    return data["candidates"][0]["content"]["parts"][0]["text"].strip()


def _extract_json(text: str) -> dict:
    """Extract the first JSON object from Gemini output."""
    start = text.find("{")
    end = text.rfind("}") + 1
    if start == -1 or end == 0:
        raise ValueError(f"JSON not found in Gemini output:\n{text[:300]}")
    return json.loads(text[start:end])


# ── Spec generation ────────────────────────────────────────────────────────────

_SPEC_PROMPT = """\
あなたは TextSeries の製品設計AIです。
以下の情報から、新製品の完全な仕様をJSONで返してください。

製品スラッグ: {slug}
ターゲット業界: {industry}
書類種別リスト: {docs}
絵文字: {emoji}
製品名: {product_name}

以下の形式のJSONを返してください（コードブロック不要、JSONのみ）:

{{
  "tagline": "書類名1・書類名2をAIで30秒に",
  "subtitle": "{industry}向けAI書類ジェネレーター",
  "description": "ログイン画面の1〜2文説明（例: 顧客情報を入力するだけで、見積書・提案書のドラフトを自動生成します。）",
  "chatgpt_diff": [
    "**業務特化プロンプト**: ...",
    "**履歴が蓄積される**: ...",
    "**Word形式で即ダウンロード**: ..."
  ],
  "company_label": "顧問先の会社名（例: 事業所名、施主名、会社名）",
  "company_placeholder": "例: ...",
  "gyoshu_label": "業種（例: 業種・施設種別・工事種別など）",
  "gyoshu_placeholder": "例: ...",
  "employees_label": "規模（例: 従業員数、施設規模など）",
  "employees_placeholder": "例: ...",
  "form_fields": {{
    "書類種別名": [
      {{"name": "field_name", "label": "フィールド名", "placeholder": "例: ...", "type": "text", "optional": false}},
      {{"name": "field2", "label": "フィールド2", "placeholder": "例: ...", "type": "textarea", "height": 100, "optional": false}},
      {{"name": "field3", "label": "任意フィールド", "placeholder": "例: ...", "type": "text", "optional": true}}
    ]
  }},
  "prompts": {{
    "書類種別名": "あなたは{industry}のアシスタントとして、書類名のドラフトを作成します。\\n\\n以下の情報をもとに...\\n\\n【案件情報】\\n{{company_label}}: {{company}}\\n{{gyoshu_label}}: {{gyoshu}}\\n{{employees_label}}: {{employees}}\\nfield_label: {{field_name}}\\n...\\n\\n【出力ルール】\\n1. ...\\n2. ...\\n3. ...\\n4. {{MIN_CHARS}}字以上で作成すること\\n5. 末尾に「本書類はドラフトです。専門家確認が必要です。」と注記\\n\\nドラフト本文のみ出力してください:"
  }},
  "icp": {{
    "target_description": "ターゲット担当者の説明",
    "company_name_hint_keywords": ["キーワード1", "キーワード2"],
    "target_keywords": ["検索キーワード1", "検索キーワード2"],
    "good_domain_suffixes": ["co.jp", "jp"],
    "bad_domain_suffixes": ["gmail.com", "yahoo.co.jp", "hotmail.com"],
    "exclude_keywords": ["求人", "転職", "ランキング", "比較", "まとめ", "選び方", "一覧"]
  }},
  "lp": {{
    "hero_pain": "書類、まだ毎回ゼロから作ってますか？",
    "hero_description": "情報を入力するだけで30秒で文案生成。",
    "pain_points": [
      {{"icon": "⏰", "title": "タイトル", "desc": "説明文"}},
      {{"icon": "📋", "title": "タイトル", "desc": "説明文"}},
      {{"icon": "😓", "title": "タイトル", "desc": "説明文"}}
    ],
    "features": [
      {{"icon": "📂", "title": "{docs_count}書類種別に対応", "desc": "説明文"}},
      {{"icon": "🏢", "title": "業務に即した文案", "desc": "説明文"}},
      {{"icon": "⚖️", "title": "法令・業界標準を踏まえた設計", "desc": "説明文"}},
      {{"icon": "💾", "title": "情報を保存・再利用", "desc": "説明文"}}
    ],
    "docs_categories": [
      {{"cat": "カテゴリ名", "items": ["書類名1", "書類名2"]}}
    ],
    "faq": [
      {{"q": "生成した書類をそのまま使えますか？", "a": "担当者による確認・修正が必要です。..."}},
      {{"q": "入力した情報は安全ですか？", "a": "入力した情報はAI生成のためにのみ使用します。..."}}
    ],
    "color_hex": "#1a4a7a",
    "color_light_hex": "#dbeafe",
    "tailwind_color": "blue"
  }},
  "email": {{
    "subject": "【書類作成をAIで効率化】{product_name}のご紹介",
    "fallback_opening": "突然のご連絡をご容赦ください。{industry}の実務でご多忙のことと存じます。",
    "personalize_prompt": "{industry}「{{company_name}}」の担当者に向けた、書類作成業務への共感と課題提起の一文を書いてください。「御社では〜」で始め、2文以内で。",
    "sequence_1": "（初回メール本文: 200〜300字、書類作成の課題を提起して{product_name}を紹介、無料トライアルへ誘導）",
    "sequence_2": "（フォローアップメール本文: 150〜200字、前回のメールへの言及と{product_name}の具体的なメリット強調）"
  }},
  "feedback_prompt_context": "{product_name}は{industry}向けAI書類文案生成ツール。ユーザーは{target}。{docs_str}の文案をAIが自動生成する。",
  "outreach": {{
    "search_queries": [
      "{industry} \"お問い合わせ\" site:co.jp -求人",
      "{industry} \"info@\" -採用",
      "{industry} \"メールでのお問い合わせ\" -転職",
      "（計10〜15件。Brave API で実際にヒットしそうな具体的なクエリを業種に合わせて生成すること）"
    ],
    "company_required_keywords": ["{industry}に特徴的な会社名キーワード（例: 特別養護 老人ホーム）"],
    "blog_signals": ["コツ", "選び方", "ランキング", "比較", "まとめ", "一覧", "営業リスト"]
  }}
}}

重要:
- prompts内の{{...}}はPython str.formatのプレースホルダー。form_fieldsのnameと完全一致させること
- 必ず {{MIN_CHARS}} を各プロンプト末尾付近に含めること
- company/gyoshu/employeesは全書類種別の共通フィールド（prompts内で必ず使うこと）
- form_fieldsのoptional=trueのフィールドはexpander内に配置される
- tailwind_colorはtailwindのカラー名（blue/green/orange/purple等）
- search_queriesは実際に{industry}のメールアドレス付きサイトがヒットしそうな具体的なクエリを15件生成すること
"""


def generate_spec(slug: str, industry: str, docs: list[str], emoji: str,
                  product_name: str, api_key: str) -> dict:
    print("  Gemini でプロダクト仕様を生成中...")
    prompt = _SPEC_PROMPT.format(
        slug=slug,
        industry=industry,
        docs=", ".join(docs),
        emoji=emoji,
        product_name=product_name,
        docs_count=len(docs),
        target=f"{industry}の担当者・管理者",
        docs_str="・".join(docs),
    )
    for attempt in range(3):
        try:
            raw = _gemini(prompt, api_key, max_tokens=16000)
            spec = _extract_json(raw)
            # Validate minimum required fields
            for required in ("tagline", "prompts", "form_fields", "icp"):
                if required not in spec:
                    raise ValueError(f"Missing field: {required}")
            return spec
        except (ValueError, KeyError, json.JSONDecodeError) as e:
            if attempt == 2:
                raise RuntimeError(f"Gemini spec 生成失敗（3回試行）: {e}") from e
            print(f"  リトライ ({attempt + 1}/3): {e}")
    raise RuntimeError("unreachable")


# ── File writers ───────────────────────────────────────────────────────────────

def _write(path: Path, content: str, dry_run: bool) -> None:
    if dry_run:
        print(f"    [dry-run] write {path}")
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"    ✓ {path.relative_to(path.anchor)}")


def write_requirements_txt(dest: Path, dry_run: bool) -> None:
    content = (
        "streamlit>=1.32.0\n"
        "python-dotenv>=1.0.0\n"
        "python-docx>=1.1.0\n"
        "supabase>=2.3.0\n"
    )
    _write(dest / "requirements.txt", content, dry_run)


def write_db_py(slug: str, dest: Path, dry_run: bool) -> None:
    content = f'''"""Supabase REST APIを直接呼び出してトライアル使用量とコードを管理する。"""
from __future__ import annotations
import os
import urllib.request
import urllib.error
import urllib.parse
import json
from datetime import datetime, timezone, timedelta

T_TRIALS = "{slug}_trials"
T_CODES = "{slug}_codes"
T_HISTORY = "{slug}_history"
T_FEEDBACK = "{slug}_feedback"


def _headers() -> dict:
    key = os.environ.get("SUPABASE_SERVICE_KEY") or os.environ["SUPABASE_ANON_KEY"]
    key = key.strip()
    return {{
        "apikey": key,
        "Authorization": f"Bearer {{key}}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }}


def _url(table: str, query: str = "") -> str:
    base = os.environ["SUPABASE_URL"].strip().rstrip("/")
    if base.endswith("/rest/v1"):
        base = base[:-len("/rest/v1")]
    url = f"{{base}}/rest/v1/{{table}}"
    if query:
        url += f"?{{query}}"
    return url


def _get(table: str, query: str) -> list:
    req = urllib.request.Request(_url(table, query), headers=_headers())
    with urllib.request.urlopen(req, timeout=10) as res:
        return json.loads(res.read())


def _post(table: str, data: dict) -> dict:
    payload = json.dumps(data).encode()
    req = urllib.request.Request(_url(table), data=payload, headers=_headers(), method="POST")
    with urllib.request.urlopen(req, timeout=10) as res:
        result = json.loads(res.read())
        return result[0] if isinstance(result, list) else result


def _patch(table: str, query: str, data: dict) -> None:
    payload = json.dumps(data).encode()
    req = urllib.request.Request(_url(table, query), data=payload, headers=_headers(), method="PATCH")
    with urllib.request.urlopen(req, timeout=10) as res:
        res.read()


def get_or_create_user(email: str) -> dict:
    try:
        rows = _get(T_TRIALS, f"email=eq.{{urllib.parse.quote(email)}}&select=*")
    except urllib.error.HTTPError as e:
        body = e.read().decode()[:200]
        raise RuntimeError(f"GET trials → {{e.code}}: {{body}}") from None
    if rows:
        return rows[0]
    try:
        return _post(T_TRIALS, {{"email": email, "count": 0, "plan": None}})
    except urllib.error.HTTPError as e:
        body = e.read().decode()[:200]
        raise RuntimeError(f"POST trials → {{e.code}}: {{body}}") from None


def increment_count(email: str) -> int:
    user = get_or_create_user(email)
    new_count = user["count"] + 1
    _patch(T_TRIALS, f"email=eq.{{urllib.parse.quote(email)}}", {{"count": new_count}})
    return new_count


def set_plan(email: str, plan: str) -> None:
    rows = _get(T_TRIALS, f"email=eq.{{urllib.parse.quote(email)}}&select=email")
    if rows:
        _patch(T_TRIALS, f"email=eq.{{urllib.parse.quote(email)}}", {{"count": 0, "plan": plan}})
    else:
        _post(T_TRIALS, {{"email": email, "count": 0, "plan": plan}})


def validate_code(code: str) -> str | None:
    rows = _get(T_CODES, f"code=eq.{{urllib.parse.quote(code)}}&active=eq.true&select=plan")
    return rows[0]["plan"] if rows else None


def save_generation(email: str, params: dict, result: str) -> None:
    try:
        _post(T_HISTORY, {{
            "email": email,
            "doc_type": params.get("doc_type", ""),
            "company": params.get("company", ""),
            "gyoshu": params.get("gyoshu", ""),
            "draft": result,
        }})
    except Exception:
        pass


def save_feedback(email: str, category: str, rating: str, regen_count: int = 0, reasons: str = "") -> None:
    try:
        _post(T_FEEDBACK, {{
            "email": email,
            "category": category,
            "rating": rating,
            "regen_count": regen_count,
            "reasons": reasons,
        }})
    except Exception:
        pass


def get_history(email: str, limit: int = 10) -> list[dict]:
    try:
        return _get(
            T_HISTORY,
            f"email=eq.{{urllib.parse.quote(email)}}"
            f"&order=created_at.desc&limit={{limit}}&select=*",
        )
    except Exception:
        return []


def get_stats(email: str) -> dict[str, int]:
    now = datetime.now(timezone.utc)
    week_start = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    try:
        week_rows = _get(T_HISTORY, f"email=eq.{{urllib.parse.quote(email)}}&created_at=gte.{{urllib.parse.quote(week_start.isoformat())}}&select=id")
        month_rows = _get(T_HISTORY, f"email=eq.{{urllib.parse.quote(email)}}&created_at=gte.{{urllib.parse.quote(month_start.isoformat())}}&select=id")
        return {{"week": len(week_rows), "month": len(month_rows)}}
    except Exception:
        return {{"week": 0, "month": 0}}
'''
    _write(dest / "db.py", content, dry_run)


def _build_form_fields_code(doc_type: str, fields: list[dict],
                             optional_label: str = "特記事項（任意）") -> tuple[str, list[str]]:
    """Generate Python form field code and params list for one doc type."""
    lines = []
    param_names = ["company", "gyoshu", "employees"]
    optional_fields = [f for f in fields if f.get("optional")]
    required_fields = [f for f in fields if not f.get("optional")]

    for field in required_fields:
        name = field["name"]
        label = field["label"]
        placeholder = field.get("placeholder", "")
        ftype = field.get("type", "text")
        if ftype == "selectbox":
            options = field.get("options", [])
            lines.append(f'        {name} = st.selectbox("{label}", {options!r})')
        elif ftype in ("textarea", "text_area"):
            height = field.get("height", 100)
            lines.append(f'        {name} = st.text_area("{label}", placeholder="{placeholder}", height={height})')
        else:
            lines.append(f'        {name} = st.text_input("{label}", placeholder="{placeholder}")')
        param_names.append(name)

    if optional_fields:
        lines.append(f'        with st.expander("{optional_label}（任意）"):')
        for field in optional_fields:
            name = field["name"]
            label = field["label"]
            placeholder = field.get("placeholder", "")
            ftype = field.get("type", "text")
            height = field.get("height", 80)
            if ftype in ("textarea", "text_area"):
                lines.append(f'            {name} = st.text_area("", placeholder="{placeholder}", height={height}, label_visibility="collapsed")')
            else:
                lines.append(f'            {name} = st.text_input("", placeholder="{placeholder}", label_visibility="collapsed")')
            param_names.append(name)

    params_dict = "{" + ", ".join(
        f'"{k}": {k}' + (' or "なし"' if any(f["name"] == k and f.get("optional") for f in fields) else "")
        for k in param_names
    ) + "}"
    lines.append(f'        params = lambda: {params_dict}')

    return "\n".join(lines), param_names


def write_app_py(slug: str, product_name: str, emoji: str, spec: dict,
                 docs: list[str], dest: Path, dry_run: bool) -> None:
    slug_upper = slug.upper()
    company_label = spec.get("company_label", "会社名・事業所名")
    company_placeholder = spec.get("company_placeholder", "例: 株式会社〇〇")
    gyoshu_label = spec.get("gyoshu_label", "業種")
    gyoshu_placeholder = spec.get("gyoshu_placeholder", "例: 製造業")
    employees_label = spec.get("employees_label", "従業員数・規模")
    employees_placeholder = spec.get("employees_placeholder", "例: 20名")
    subtitle = spec.get("subtitle", f"{product_name}")
    description = spec.get("description", "情報を入力するだけで書類のドラフトを自動生成します。")
    chatgpt_diff = spec.get("chatgpt_diff", [
        "**業務特化プロンプト**: 業界標準を熟知したAIが生成",
        "**履歴が蓄積される**: 過去の書類が自動保存",
        "**Word形式で即ダウンロード**: .docx形式で即使える",
    ])
    chatgpt_diff_md = "\n".join(f"    - {item}" for item in chatgpt_diff)

    # Build form body
    form_fields_spec = spec.get("form_fields", {})
    form_blocks = []
    for dt in docs:
        fields = form_fields_spec.get(dt, [])
        field_code, _ = _build_form_fields_code(dt, fields)
        is_first = dt == docs[0]
        is_last = dt == docs[-1]
        if is_first:
            form_blocks.append(f'    if doc_type == "{dt}":')
        elif is_last:
            form_blocks.append(f'    else:  # {dt}')
        else:
            form_blocks.append(f'    elif doc_type == "{dt}":')
        form_blocks.append(field_code)
    form_body = "\n\n".join(form_blocks)

    doc_types_repr = repr(docs)
    tos_description = spec.get("subtitle", f"{product_name}")
    page_title_str = "・".join(docs[:2]) + " AI | " + product_name

    content = f'''import os
import streamlit as st

STRIPE_STANDARD_URL = st.secrets.get("{slug_upper}_STRIPE_STANDARD_URL", "")
STRIPE_PRO_URL = st.secrets.get("{slug_upper}_STRIPE_PRO_URL", "")
CONTACT_EMAIL = "ryuumg03@gmail.com"
FREE_TRIAL_LIMIT = 5
PLAN_LIMITS = {{"standard": 50, "pro": 200}}

if "GEMINI_API_KEY" in st.secrets:
    os.environ["GEMINI_API_KEY"] = st.secrets["GEMINI_API_KEY"]
    os.environ["SUPABASE_URL"] = st.secrets.get("SUPABASE_URL", "")
    os.environ["SUPABASE_ANON_KEY"] = st.secrets.get("SUPABASE_ANON_KEY", "")
else:
    try:
        from dotenv import load_dotenv
        load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), "../../../../.env"))
    except ImportError:
        pass

_USE_DB = bool(os.environ.get("SUPABASE_URL") and os.environ.get("SUPABASE_ANON_KEY"))
if _USE_DB:
    from db import get_or_create_user, increment_count, set_plan, validate_code, save_generation, get_history, get_stats, save_feedback

st.set_page_config(page_title="{page_title_str}", page_icon="{emoji}", layout="centered")


with st.expander("ChatGPTとの違い・{product_name}を使う理由", expanded=False):
    st.markdown("""
{chatgpt_diff_md}
    """)
for key, default in [
    ("user_email", None),
    ("db_loaded", False),
    ("request_count", 0),
    ("paid_plan", None),
    ("last_result", None),
    ("last_doc_type", None),
    ("last_company", None),
    ("last_docx", None),
    ("generation_stats", {{"week": 0, "month": 0}}),
    ("feedback_sent", False),
    ("show_feedback_form", False),
]:
    if key not in st.session_state:
        st.session_state[key] = default

if not st.session_state.user_email:
    st.title("{emoji} {product_name}")
    st.markdown("**{subtitle}**  \\n{description}")
    st.markdown("---")
    email = st.text_input("メールアドレスを入力してスタート", placeholder="example@office.co.jp")
    if st.button("スタート", type="primary"):
        if "@" not in email or "." not in email:
            st.error("有効なメールアドレスを入力してください。")
        else:
            st.session_state.user_email = email.strip().lower()
            st.rerun()
    st.stop()

if _USE_DB and not st.session_state.db_loaded:
    try:
        user = get_or_create_user(st.session_state.user_email)
        st.session_state.request_count = user.get("count", 0)
        st.session_state.paid_plan = user.get("plan")
        st.session_state.generation_stats = get_stats(st.session_state.user_email)
        st.session_state.db_loaded = True
    except Exception as e:
        st.warning(f"DB接続エラー: {{e}}")

_plan = st.session_state.paid_plan
_limit = PLAN_LIMITS.get(_plan, FREE_TRIAL_LIMIT) if _plan else FREE_TRIAL_LIMIT
remaining = _limit - st.session_state.request_count

st.title("{emoji} {product_name} — {subtitle}")

if _plan:
    plan_label = "スタンダード" if _plan == "standard" else "プロ"
    st.success(f"{{plan_label}}プラン: **{{max(0, remaining)}} / {{_limit}}件** 残り（今月）")
elif remaining > 0:
    if remaining <= 2:
        st.warning(f"無料トライアル残り **{{remaining}}件** — 続けて使うには有料プランをご検討ください。")
    else:
        st.info(f"無料トライアル: **{{remaining}} / {{FREE_TRIAL_LIMIT}}件** 残り")

if st.session_state.request_count >= _limit:
    if _plan:
        st.error(f"今月の上限（{{_limit}}件）に達しました。プランアップグレードをご検討ください。")
    else:
        st.error(f"無料トライアル（{{FREE_TRIAL_LIMIT}}件）を使い切りました。")
    st.markdown("---")
    st.markdown("### {emoji} 有料プランのご案内")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**スタンダードプラン**\\n- 月50件まで生成\\n- **¥8,980/月**")
        if STRIPE_STANDARD_URL:
            st.link_button("スタンダードに申し込む", STRIPE_STANDARD_URL, type="primary", use_container_width=True)
    with col2:
        st.markdown("**プロプラン**\\n- 月200件まで生成\\n- **¥19,800/月**")
        if STRIPE_PRO_URL:
            st.link_button("プロプランに申し込む", STRIPE_PRO_URL, use_container_width=True)
    st.caption("※ お申し込み後、アクセスコードをメールでお送りします。")
    st.markdown("---")
    st.markdown("#### すでにアクセスコードをお持ちの方")
    code_input = st.text_input("アクセスコードを入力", placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx", label_visibility="collapsed")
    if st.button("コードで解除する", type="secondary"):
        c = code_input.strip()
        new_plan = validate_code(c) if _USE_DB else None
        if new_plan:
            st.session_state.paid_plan = new_plan
            st.session_state.request_count = 0
            set_plan(st.session_state.user_email, new_plan)
            st.success(f"{{'スタンダード' if new_plan == 'standard' else 'プロ'}}プランが有効化されました！")
            st.rerun()
        else:
            st.error("コードが無効です。メールに記載のコードをご確認ください。")
    st.stop()

st.markdown("---")

DOC_TYPES = {doc_types_repr}

with st.form("{slug}_form"):
    doc_type = st.radio("書類種別", DOC_TYPES, horizontal=True)
    st.markdown("---")

    company = st.text_input("{company_label}", placeholder="{company_placeholder}")
    gyoshu = st.text_input("{gyoshu_label}", placeholder="{gyoshu_placeholder}")
    employees = st.text_input("{employees_label}", placeholder="{employees_placeholder}")

{form_body}

    submitted = st.form_submit_button("ドラフトを生成する", type="primary", use_container_width=True)

if submitted:
    if not company or not gyoshu or not employees:
        st.error("{company_label}・{gyoshu_label}・{employees_label}は必須です。")
    else:
        with st.spinner("ドラフトを生成中..."):
            from gen import generate
            p = params()
            p["doc_type"] = doc_type
            result = generate(doc_type, p)

        if result["ok"]:
            if _USE_DB:
                try:
                    increment_count(st.session_state.user_email)
                    save_generation(st.session_state.user_email, p, result["draft"])
                    st.session_state.request_count += 1
                except Exception as db_err:
                    st.warning(f"利用回数の記録に失敗しました（生成結果は表示されます）: {{db_err}}")
            st.session_state.last_result = result["draft"]
            st.session_state.last_doc_type = doc_type
            st.session_state.last_company = company
            st.session_state.feedback_sent = False
            st.session_state.show_feedback_form = False
            from gen import to_docx
            st.session_state.last_docx = to_docx(result["draft"], doc_type, company)
        else:
            st.error(result["error"])

if st.session_state.last_result:
    st.markdown("---")
    st.markdown(f"### {{st.session_state.last_doc_type}} ドラフト")

    if st.session_state.last_docx:
        company_safe = (st.session_state.last_company or "draft").replace(" ", "_").replace("/", "_")[:20]
        filename = f"{{st.session_state.last_doc_type}}_{{company_safe}}.docx"
        st.download_button(
            label="Wordファイル (.docx) をダウンロード",
            data=st.session_state.last_docx,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )

    st.text_area("テキストで確認・コピー", st.session_state.last_result, height=400, label_visibility="collapsed")
    st.caption("※ 生成内容はドラフトです。実際の届出・運用前に専門家にご確認ください。")

    st.markdown("---")
    if not st.session_state.feedback_sent:
        st.markdown("**この書類は役に立ちましたか？**")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("👍 役立った", use_container_width=True, key="fb_good"):
                if _USE_DB:
                    save_feedback(st.session_state.user_email, st.session_state.last_doc_type, "good", st.session_state.request_count)
                st.session_state.feedback_sent = True
        with col2:
            if st.button("👎 改善が必要", use_container_width=True, key="fb_bad"):
                st.session_state.show_feedback_form = True
        if st.session_state.show_feedback_form:
            reasons = st.text_area("改善点を教えてください（任意）", key="fb_reasons", placeholder="どの点が不十分でしたか？")
            if st.button("送信", key="fb_submit"):
                if _USE_DB:
                    save_feedback(st.session_state.user_email, st.session_state.last_doc_type, "bad", st.session_state.request_count, reasons)
                st.session_state.feedback_sent = True
                st.session_state.show_feedback_form = False
    else:
        st.caption("フィードバックありがとうございます！")

if _USE_DB and st.session_state.get("generation_stats", {{}}).get("month", 0) > 0:
    stats = st.session_state.generation_stats
    st.markdown("---")
    st.markdown(f"今週 **{{stats['week']}}件** / 今月 **{{stats['month']}}件** 生成済み")
    with st.expander("過去の生成履歴"):
        history = get_history(st.session_state.user_email)
        for h in history:
            st.markdown(f"**{{h.get('doc_type', '')}}** — {{h.get('company', '')}} （{{h.get('created_at', '')[:10]}}）")
            st.text_area("", h.get("draft", ""), height=150, key=f"h_{{h.get('id', '')}}", label_visibility="collapsed")

st.markdown("---")

with st.expander("利用規約"):
    st.markdown("""
**第1条（適用）** 本規約は {product_name} の利用条件を定めます。ご利用をもって同意とみなします。

**第2条（サービスの内容）** 本サービスは{tos_description}の文書草案を生成するAIツールです。生成内容は参考資料であり、実際の届出・運用前に専門家による確認が必要です。

**第3条（禁止事項）** 虚偽情報の入力・不正アクセス・第三者への再販・第三者の情報を無断で入力する行為等を禁止します。

**第4条（免責事項）** 生成文章の正確性・完全性・法的適合性を保証しません。生成結果の使用による損害について運営者は責任を負いません。

**第5条（知的財産）** システム・デザインの権利は運営者に帰属します。生成文章の著作権は利用者に帰属します。

**第6条（個人情報）** 利用回数管理のためメールアドレスを収集します。

**第7条（準拠法）** 本規約は日本法に準拠します。紛争は運営者所在地の管轄裁判所を第一審とします。
""")

with st.expander("プライバシーポリシー"):
    st.markdown("""
**収集する情報:** メールアドレス（利用回数管理のため）。入力された情報はAI生成のみに使用します。

**保存期間:** メールアドレスおよび利用回数はサービス利用中のみ保持します。生成履歴は最大90日間保持後に自動削除します。

**収集しない情報:** 氏名・住所・電話番号等、IPログ、Cookie追跡は行いません。

**第三者提供:** 文章生成のためGoogle Gemini API（Google LLC）に入力情報を送信します。利用回数管理のためSupabase（Supabase Inc.）にメールアドレスを保存します。それ以外の第三者提供はありません。

**お問い合わせ:** ryuumg03@gmail.com
""")

with st.expander("特定商取引法に基づく表記"):
    st.markdown("""
| 項目 | 内容 |
|------|------|
| 販売事業者 | TextSeries（運営者: 真柄 龍聖） |
| 所在地 | 〒060-0001 北海道札幌市中央区北一条西3丁目3番地33 リープロビル302 |
| メールアドレス | ryuumg03@gmail.com |
| 販売価格 | スタンダード ¥8,980/月 / プロ ¥19,800/月（無料トライアル: 5件まで） |
| 支払方法 | クレジットカード（Stripe決済） |
| 提供時期 | 決済完了後、アクセスコードをメールにてご案内（通常1営業日以内） |
| 返品・キャンセル | デジタルサービスのため提供済み期間の返金不可。翌月更新日前日までの解約で翌月以降の請求なし。 |
""")

st.caption("© 2026 {product_name}　📧 ryuumg03@gmail.com")
'''
    _write(dest / "app.py", content, dry_run)


def write_gen_py(slug: str, product_name: str, spec: dict,
                 docs: list[str], dest: Path, dry_run: bool) -> None:
    prompts_spec = spec.get("prompts", {})
    prompts_lines = []
    for dt in docs:
        raw_prompt = prompts_spec.get(dt, f"あなたは{product_name}のアシスタントとして{dt}のドラフトを作成します。\n\n【情報】\n対象: {{company}}\n業種: {{gyoshu}}\n規模: {{employees}}\n\n{{MIN_CHARS}}字以上で作成してください。\n\nドラフト本文のみ出力してください:")
        escaped = raw_prompt.replace('"""', '\\"\\"\\"').replace("\\n", "\n")
        prompts_lines.append(f'    {json.dumps(dt, ensure_ascii=False)}: """{escaped}""",')

    prompts_block = "\n".join(prompts_lines)

    content = f'''"""Gemini APIで{product_name}向け書類ドラフトを生成する。"""
import io
import os
import json
import re
import urllib.request
import urllib.error
from datetime import date

MODEL = "gemini-2.5-flash"
MAX_TOKENS = 1800
MIN_CHARS = 400
MAX_RETRIES = 3

_PROMPTS = {{
{prompts_block}
}}

_INJECTION_PATTERNS = [
    "ignore previous", "前の指示を無視", "system:", "【出力ルール】", "---\\n",
]


def _sanitize(text: str, max_len: int = 200) -> str:
    text = text.replace("\\n", " ").replace("\\t", " ")
    while "  " in text:
        text = text.replace("  ", " ")
    for p in _INJECTION_PATTERNS:
        if p.lower() in text.lower():
            return ""
    return text[:max_len].strip()


def generate(doc_type: str, params: dict) -> dict:
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key:
        return {{"ok": False, "error": "GEMINI_API_KEY が設定されていません", "error_type": "config"}}

    template = _PROMPTS.get(doc_type)
    if not template:
        return {{"ok": False, "error": f"書類種別「{{doc_type}}」は未対応です", "error_type": "config"}}

    safe = {{k: _sanitize(str(v)) for k, v in params.items()}}
    safe["MIN_CHARS"] = MIN_CHARS

    try:
        prompt = template.format(**safe)
    except KeyError as e:
        return {{"ok": False, "error": f"パラメータ不足: {{e}}", "error_type": "config"}}

    payload = json.dumps({{
        "contents": [{{"parts": [{{"text": prompt}}]}}],
        "generationConfig": {{"maxOutputTokens": MAX_TOKENS, "temperature": 0.3}},
    }}).encode()

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{{MODEL}}:generateContent?key={{api_key}}"

    for attempt in range(MAX_RETRIES):
        try:
            req = urllib.request.Request(url, data=payload, headers={{"Content-Type": "application/json"}}, method="POST")
            with urllib.request.urlopen(req, timeout=30) as r:
                data = json.loads(r.read())
            text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
            if len(text) < MIN_CHARS:
                continue
            return {{"ok": True, "draft": text}}
        except urllib.error.HTTPError as e:
            err = e.read().decode()
            if "429" in err or "RESOURCE_EXHAUSTED" in err:
                return {{"ok": False, "error": "APIの利用上限に達しました。しばらくしてから再試行してください。", "error_type": "quota"}}
            if attempt == MAX_RETRIES - 1:
                return {{"ok": False, "error": f"生成に失敗しました（{{e.code}}）", "error_type": "api"}}
        except Exception:
            if attempt == MAX_RETRIES - 1:
                return {{"ok": False, "error": "生成中にエラーが発生しました", "error_type": "unknown"}}

    return {{"ok": False, "error": "生成に失敗しました。もう一度お試しください。", "error_type": "retry"}}


_ARTICLE_RE = re.compile(r"^(第\\d+条[^\\s　]{{0,20}})", re.MULTILINE)
_CHAPTER_RE = re.compile(r"^(第\\d+章[^\\s　]{{0,20}})", re.MULTILINE)


def to_docx(draft: str, doc_type: str, company: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(doc_type, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(company).font.size = Pt(11)
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"生成日: {{date.today().strftime('%Y年%m月%d日')}}").font.size = Pt(9)
    doc.add_paragraph()

    lines = draft.split("\\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if _CHAPTER_RE.match(stripped):
            h = doc.add_heading(stripped, level=1)
            h.runs[0].font.size = Pt(13)
        elif _ARTICLE_RE.match(stripped):
            h = doc.add_heading(stripped, level=2)
            h.runs[0].font.size = Pt(11)
        elif stripped.startswith(("　", "  ", "1.", "2.", "3.", "・", "※", "【")):
            p = doc.add_paragraph(stripped)
            p.runs[0].font.size = Pt(10.5)
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            p = doc.add_paragraph(stripped)
            p.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()
    note = doc.add_paragraph("※ 本書類はAIが生成したドラフトです。実際の運用・届出前に専門家にご確認ください。")
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
'''
    _write(dest / "gen.py", content, dry_run)


def write_config_json(slug: str, product_name: str, rank: int,
                      spec: dict, docs: list[str],
                      dest: Path, dry_run: bool) -> None:
    icp_spec = spec.get("icp", {})
    email_spec = spec.get("email", {})
    config = {
        "project": slug,
        "project_dir": slug,
        "product_name": product_name,
        "sender_name": "TextSeries",
        "sender_email": "ryuumg03@gmail.com",
        "app_url": f"https://{slug}.streamlit.app",
        "lp_url": f"https://ryuu321.github.io/ai-holdings/docs/{slug}.html",
        "daily_send_limit": 30,
        "send_interval_sec": 30,
        "gemini_model": "gemini-2.5-flash",
        "icp": {
            "target_description": icp_spec.get("target_description", ""),
            "document_types": docs,
            "company_name_hint_keywords": icp_spec.get("company_name_hint_keywords", []),
            "target_keywords": icp_spec.get("target_keywords", []),
            "good_domain_suffixes": icp_spec.get("good_domain_suffixes", ["co.jp", "jp"]),
            "bad_domain_suffixes": icp_spec.get("bad_domain_suffixes", ["gmail.com", "yahoo.co.jp"]),
            "exclude_keywords": icp_spec.get("exclude_keywords", ["求人", "転職", "ランキング"]),
        },
        "scoring": {
            "target_keyword_hit": 30,
            "good_domain": 20,
            "has_company_type": 20,
            "bad_domain_penalty": -20,
            "exclude_keyword_penalty": -50,
            "auto_approve_threshold": 70,
            "review_threshold": 50,
        },
        "sender_address": "〒060-0001 北海道札幌市中央区北一条西3丁目3番地33 リープロビル302",
        "sent_log": f"saas-dev/projects/{slug}/outreach/sent_log.csv",
        "sent_log_gist_filename": f"{slug}_sent_log.csv",
        "followup_days": 7,
        "followup_template": f"shared/gtm/outreach/templates/{slug}_sequence_2.txt",
        "feedback_csv_env": "",
        "feedback_report_dir": f"saas-dev/projects/{slug}/feedback_reports",
        "feedback_prompt_context": spec.get(
            "feedback_prompt_context",
            f"{product_name}は{icp_spec.get('target_description', '')}向けAI書類文案生成ツール。"
        ),
        "stripe_standard_url": "",
        "stripe_pro_url": "",
        "history_columns": [
            "doc_type text",
            "company text",
            "gyoshu text",
            "draft text",
        ],
        "email_template": {
            "template_file": f"{slug}_sequence_1.txt",
            "subject": email_spec.get("subject", f"【書類作成をAIで効率化】{product_name}のご紹介"),
            "fallback_opening": email_spec.get("fallback_opening", "突然のご連絡をご容赦ください。"),
            "personalize_prompt": email_spec.get("personalize_prompt", ""),
        },
    }
    _write(
        dest / "shared" / "gtm" / "config" / f"{slug}.json",
        json.dumps(config, ensure_ascii=False, indent=2),
        dry_run,
    )


def write_lp_html(slug: str, product_name: str, emoji: str,
                  spec: dict, docs: list[str], dest: Path, dry_run: bool) -> None:
    lp = spec.get("lp", {})
    tc = lp.get("tailwind_color", "blue")
    color_hex = lp.get("color_hex", "#1a4a7a")
    color_light = lp.get("color_light_hex", "#dbeafe")

    tagline = spec.get("tagline", f"{'・'.join(docs[:2])}をAIで30秒に")
    hero_pain = lp.get("hero_pain", f"{docs[0]}、まだ毎回ゼロから作ってますか？")
    hero_desc = lp.get("hero_description", "情報を入力するだけで30秒で文案生成。")

    def pain_html(pts: list) -> str:
        cards = ""
        for pt in pts[:3]:
            cards += f'''
        <div class="bg-white rounded-xl p-6 shadow-sm border border-gray-100">
          <div class="text-3xl mb-3">{pt.get("icon", "❓")}</div>
          <p class="font-semibold mb-1">{pt.get("title", "")}</p>
          <p class="text-gray-500 text-sm">{pt.get("desc", "")}</p>
        </div>'''
        return cards

    def feature_html(features: list) -> str:
        cards = ""
        for f in features[:4]:
            cards += f'''
        <div class="feature-card bg-{tc}-50 rounded-xl p-6 border border-{tc}-100">
          <div class="text-2xl mb-3">{f.get("icon", "✨")}</div>
          <h3 class="font-bold text-lg mb-2">{f.get("title", "")}</h3>
          <p class="text-gray-600 text-sm">{f.get("desc", "")}</p>
        </div>'''
        return cards

    def docs_cat_html(cats: list) -> str:
        cards = ""
        for cat in cats:
            items_html = "".join(f"<li>• {i}</li>" for i in cat.get("items", []))
            cards += f'''
        <div class="bg-gray-50 rounded-lg p-4">
          <p class="font-semibold text-{tc}-800 mb-2">{cat.get("cat", "")}</p>
          <ul class="text-gray-600 text-sm space-y-1">{items_html}</ul>
        </div>'''
        return cards

    def faq_html(faqs: list) -> str:
        items = ""
        for faq in faqs:
            items += f'''
        <details class="bg-gray-50 rounded-xl p-5 cursor-pointer">
          <summary class="font-semibold text-gray-800">{faq.get("q", "")}</summary>
          <p class="text-gray-600 text-sm mt-3">{faq.get("a", "")}</p>
        </details>'''
        return items

    pain_pts = lp.get("pain_points", [
        {"icon": "⏰", "title": "書類作成に時間がかかる", "desc": "毎回ゼロから書類を作成するため、本来の業務に集中できない"},
        {"icon": "📋", "title": "品質にばらつきがある", "desc": "担当者によって記載内容の詳しさ・正確さに差が出てしまう"},
        {"icon": "😓", "title": "書類業務を後回しにしてしまう", "desc": "忙しい中での書類作成が負担となり、後回しになりがち"},
    ])
    features = lp.get("features", [
        {"icon": "📂", "title": f"{len(docs)}書類種別に対応", "desc": "業務で必要な書類をワンストップで生成できます。"},
        {"icon": "🏢", "title": "業務に即した文案", "desc": "業界の実務を踏まえたAIが、現場で使える文案を自動生成します。"},
        {"icon": "⚖️", "title": "法令・業界標準を踏まえた設計", "desc": "関連法令・業界標準の要件に沿った文案を生成します。"},
        {"icon": "💾", "title": "情報を保存・再利用", "desc": "過去の書類・顧客情報を保存。次回からはすぐ再利用できます。"},
    ])
    docs_cats = lp.get("docs_categories", [{"cat": "書類一覧", "items": docs}])
    faqs = lp.get("faq", [
        {"q": "生成した書類をそのまま使えますか？", "a": "担当者による確認・修正が必要です。固有名詞・数値・日付を追記し、専門家確認のうえご使用ください。"},
        {"q": "入力した情報は安全ですか？", "a": "入力した情報はAI文案生成のためにのみ使用します。個人情報の入力は避けてください。"},
        {"q": "解約はいつでもできますか？", "a": "はい、いつでも解約できます。次回更新日前日までの解約で翌月以降の請求はありません。"},
    ])

    html = f'''<!DOCTYPE html>
<html lang="ja">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{product_name} — {tagline}</title>
  <meta name="description" content="{hero_desc}">
  <script src="https://cdn.tailwindcss.com"></script>
  <style>
    .gradient-hero {{ background: linear-gradient(135deg, {color_hex} 0%, {color_hex}cc 100%); }}
    .feature-card:hover {{ transform: translateY(-4px); transition: transform 0.2s; }}
  </style>
</head>
<body class="font-sans text-gray-800 bg-white">

  <header class="gradient-hero text-white py-20 px-6 text-center">
    <p class="text-{tc}-200 text-sm font-medium mb-4 tracking-widest uppercase">{spec.get("subtitle", "")}</p>
    <h1 class="text-4xl md:text-5xl font-bold leading-tight mb-6">
      {hero_pain}
    </h1>
    <p class="text-xl text-{tc}-100 mb-10 max-w-xl mx-auto">
      {hero_desc}<br>
      {'・'.join(docs)}が自動で完成。
    </p>
    <a href="https://{slug}.streamlit.app"
       target="_blank"
       class="inline-block bg-yellow-400 hover:bg-yellow-300 text-gray-900 font-bold text-lg px-10 py-4 rounded-full shadow-lg transition">
      無料で試してみる
    </a>
    <p class="text-{tc}-200 text-sm mt-4">クレジットカード不要・登録不要・5件まで無料</p>
  </header>

  <section class="py-16 px-6 bg-gray-50">
    <div class="max-w-3xl mx-auto text-center">
      <h2 class="text-2xl font-bold mb-10 text-gray-700">こんな悩み、ありませんか？</h2>
      <div class="grid md:grid-cols-3 gap-6 text-left">{pain_html(pain_pts)}
      </div>
    </div>
  </section>

  <section class="py-16 px-6">
    <div class="max-w-3xl mx-auto text-center">
      <h2 class="text-3xl font-bold mb-4">{product_name} が解決します</h2>
      <p class="text-gray-500 mb-12">業界の実務を知ったAIが、現場で使える書類の文案を自動で仕上げます</p>
      <div class="grid md:grid-cols-2 gap-6 text-left">{feature_html(features)}
      </div>
    </div>
  </section>

  <section class="py-16 px-6 bg-gray-50">
    <div class="max-w-2xl mx-auto text-center">
      <h2 class="text-2xl font-bold mb-12">使い方はシンプル</h2>
      <div class="space-y-8">
        <div class="flex items-start gap-6 text-left">
          <div class="flex-shrink-0 w-10 h-10 bg-{tc}-700 text-white rounded-full flex items-center justify-center font-bold text-lg">1</div>
          <div>
            <p class="font-semibold text-lg">書類種別・必要情報を入力</p>
            <p class="text-gray-500 text-sm">書類の種類と必要情報を入力（1〜2分）</p>
          </div>
        </div>
        <div class="flex items-start gap-6 text-left">
          <div class="flex-shrink-0 w-10 h-10 bg-{tc}-700 text-white rounded-full flex items-center justify-center font-bold text-lg">2</div>
          <div>
            <p class="font-semibold text-lg">AIが文案を自動生成（30秒）</p>
            <p class="text-gray-500 text-sm">必要な記載事項を含む文案を30秒で生成</p>
          </div>
        </div>
        <div class="flex items-start gap-6 text-left">
          <div class="flex-shrink-0 w-10 h-10 bg-{tc}-700 text-white rounded-full flex items-center justify-center font-bold text-lg">3</div>
          <div>
            <p class="font-semibold text-lg">確認・修正して完成</p>
            <p class="text-gray-500 text-sm">生成された文案を確認し、固有名詞・数値・日付を追記して完成</p>
          </div>
        </div>
      </div>
    </div>
  </section>

  <section class="py-16 px-6">
    <div class="max-w-3xl mx-auto text-center">
      <h2 class="text-2xl font-bold mb-10">対応書類一覧</h2>
      <div class="grid md:grid-cols-2 gap-4 text-left">{docs_cat_html(docs_cats)}
      </div>
    </div>
  </section>

  <section class="py-16 px-6 bg-gray-50">
    <div class="max-w-2xl mx-auto text-center">
      <h2 class="text-2xl font-bold mb-10">料金プラン</h2>
      <div class="grid md:grid-cols-3 gap-6">
        <div class="bg-white rounded-xl p-6 shadow-sm border border-gray-200">
          <p class="text-gray-500 text-sm mb-2">無料トライアル</p>
          <p class="text-3xl font-bold mb-1">¥0</p>
          <p class="text-gray-400 text-sm mb-4">5件まで</p>
          <ul class="text-gray-600 text-sm text-left space-y-1">
            <li>✓ 全{len(docs)}書類種別に対応</li>
            <li>✓ 登録不要</li>
          </ul>
        </div>
        <div class="bg-{tc}-700 text-white rounded-xl p-6 shadow-lg border border-{tc}-600">
          <p class="text-{tc}-200 text-sm mb-2">スタンダード</p>
          <p class="text-3xl font-bold mb-1">¥8,980<span class="text-lg font-normal">/月</span></p>
          <p class="text-{tc}-200 text-sm mb-4">月50件まで</p>
          <ul class="text-{tc}-100 text-sm text-left space-y-1">
            <li>✓ 全機能利用可能</li>
            <li>✓ 生成履歴の閲覧</li>
            <li>✓ 情報の保存・再利用</li>
          </ul>
        </div>
        <div class="bg-white rounded-xl p-6 shadow-sm border border-gray-200">
          <p class="text-gray-500 text-sm mb-2">プロ</p>
          <p class="text-3xl font-bold mb-1">¥19,800<span class="text-lg font-normal">/月</span></p>
          <p class="text-gray-400 text-sm mb-4">月200件まで</p>
          <ul class="text-gray-600 text-sm text-left space-y-1">
            <li>✓ スタンダードの全機能</li>
            <li>✓ 大量・高頻度利用に対応</li>
          </ul>
        </div>
      </div>
      <div class="mt-8">
        <a href="https://{slug}.streamlit.app"
           target="_blank"
           class="inline-block bg-{tc}-700 hover:bg-{tc}-600 text-white font-bold text-lg px-10 py-4 rounded-full shadow-lg transition">
          無料で試してみる
        </a>
      </div>
    </div>
  </section>

  <section class="py-16 px-6">
    <div class="max-w-2xl mx-auto">
      <h2 class="text-2xl font-bold mb-10 text-center">よくある質問</h2>
      <div class="space-y-4">{faq_html(faqs)}
      </div>
    </div>
  </section>

  <section class="py-16 px-6 gradient-hero text-white text-center">
    <h2 class="text-3xl font-bold mb-4">まずは無料でお試しください</h2>
    <p class="text-{tc}-100 mb-8">クレジットカード不要・登録不要・5件まで完全無料</p>
    <a href="https://{slug}.streamlit.app"
       target="_blank"
       class="inline-block bg-yellow-400 hover:bg-yellow-300 text-gray-900 font-bold text-lg px-10 py-4 rounded-full shadow-lg transition">
      {product_name}を試してみる
    </a>
  </section>

  <footer class="py-12 px-6 bg-gray-900 text-gray-400 text-sm">
    <div class="max-w-3xl mx-auto flex flex-col md:flex-row justify-between items-center gap-4">
      <div>
        <p class="font-bold text-white text-lg mb-1">{emoji} {product_name}</p>
        <p>{spec.get("subtitle", "")}</p>
      </div>
      <div class="text-right">
        <p>© 2026 TextSeries</p>
        <p class="mt-1"><a href="mailto:ryuumg03@gmail.com" class="hover:text-white">ryuumg03@gmail.com</a></p>
      </div>
    </div>
  </footer>

</body>
</html>'''
    _write(dest / "docs" / f"{slug}.html", html, dry_run)


def write_workflows(slug: str, product_name: str, dest: Path, dry_run: bool) -> None:
    slug_upper = slug.upper()

    # daily-send
    daily_send = f'''name: {slug}-daily-send

on:
  schedule:
    - cron: "30 2 * * 1-5"  # 平日 JST 11:30 (UTC 02:30)
  workflow_dispatch:
    inputs:
      dry_run:
        description: 'dry-run (true/false)'
        default: 'true'
      limit:
        description: '送信件数上限'
        default: '30'

jobs:
  send:
    runs-on: ubuntu-latest

    steps:
      - uses: actions/checkout@v4

      - uses: actions/setup-python@v5
        with:
          python-version: "3.11"

      - name: Install dependencies
        run: pip install python-dotenv

      - name: Restore sent_log from Gist
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          curl -sf -H "Authorization: token $GIST_TOKEN" \\
            "https://api.github.com/gists/$GIST_ID" \\
            | python3 -c "
          import sys, json
          data = json.load(sys.stdin)
          content = data['files'].get('{slug}_sent_log.csv', {{}}).get('content', '')
          print(content, end='')
          " > saas-dev/projects/{slug}/outreach/sent_log.csv || echo "Gist取得失敗"

      - name: Run GTM pipeline
        env:
          BRAVE_API_KEY: ${{{{ secrets.BRAVE_API_KEY }}}}
          GEMINI_API_KEY: ${{{{ secrets.GEMINI_API_KEY }}}}
        run: python saas-dev/projects/{slug}/outreach/pipeline.py --limit ${{{{ inputs.limit || '30' }}}}

      - name: Send emails
        env:
          GMAIL_ADDRESS: ${{{{ secrets.GMAIL_ADDRESS }}}}
          GMAIL_APP_PASSWORD: ${{{{ secrets.GMAIL_APP_PASSWORD }}}}
        run: |
          ARGS="--limit ${{{{ inputs.limit || '30' }}}}"
          [ "${{{{ inputs.dry_run }}}}" = "true" ] && ARGS="$ARGS --dry-run"
          python saas-dev/projects/{slug}/outreach/send_emails.py $ARGS

      - name: Upload updated sent_log to Gist
        if: always()
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          [ -f saas-dev/projects/{slug}/outreach/sent_log.csv ] || exit 0
          python3 -c "
          import json
          content = open('saas-dev/projects/{slug}/outreach/sent_log.csv', encoding='utf-8').read()
          print(json.dumps({{'files': {{'{slug}_sent_log.csv': {{'content': content}}}}}}))
          " | curl -sf -X PATCH \\
            -H "Authorization: token $GIST_TOKEN" \\
            -H "Content-Type: application/json" \\
            -d @- \\
            "https://api.github.com/gists/$GIST_ID"
'''

    # check-replies
    check_replies = f'''name: {slug}-check-replies

on:
  schedule:
    - cron: "0 1 * * 1-5"  # 平日 JST 10:00 (UTC 01:00)
  workflow_dispatch:

jobs:
  check-replies:
    runs-on: ubuntu-latest

    steps:
      - uses: actions/checkout@v4

      - uses: actions/setup-python@v5
        with:
          python-version: "3.11"

      - name: Install dependencies
        run: pip install python-dotenv

      - name: Restore sent_log from Gist
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          curl -sf -H "Authorization: token $GIST_TOKEN" \\
            "https://api.github.com/gists/$GIST_ID" \\
            | python3 -c "
          import sys, json
          data = json.load(sys.stdin)
          content = data['files'].get('{slug}_sent_log.csv', {{}}).get('content', '')
          print(content, end='')
          " > saas-dev/projects/{slug}/outreach/sent_log.csv || echo "Gist取得失敗"

      - name: Check replies and notify
        env:
          GMAIL_ADDRESS: ${{{{ secrets.GMAIL_ADDRESS }}}}
          GMAIL_APP_PASSWORD: ${{{{ secrets.GMAIL_APP_PASSWORD }}}}
          TELEGRAM_BOT_TOKEN: ${{{{ secrets.TELEGRAM_BOT_TOKEN }}}}
          TELEGRAM_CHANNEL_ID: ${{{{ secrets.TELEGRAM_CHANNEL_ID }}}}
        run: python shared/gtm/outreach/check_replies.py --project {slug} --mark

      - name: Upload updated sent_log to Gist
        if: always()
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          [ -f saas-dev/projects/{slug}/outreach/sent_log.csv ] || exit 0
          python3 -c "
          import json
          content = open('saas-dev/projects/{slug}/outreach/sent_log.csv', encoding='utf-8').read()
          print(json.dumps({{'files': {{'{slug}_sent_log.csv': {{'content': content}}}}}}))
          " | curl -sf -X PATCH \\
            -H "Authorization: token $GIST_TOKEN" \\
            -H "Content-Type: application/json" \\
            -d @- \\
            "https://api.github.com/gists/$GIST_ID"
'''

    # follow-up
    follow_up = f'''name: {slug}-follow-up

on:
  schedule:
    - cron: "30 1 * * 2"  # 毎週火曜 JST 10:30 (UTC 01:30)
  workflow_dispatch:
    inputs:
      dry_run:
        description: 'dry-run (true/false)'
        default: 'true'
      days:
        description: '何日以上前を対象にするか'
        default: '7'

jobs:
  follow-up:
    runs-on: ubuntu-latest

    steps:
      - uses: actions/checkout@v4

      - uses: actions/setup-python@v5
        with:
          python-version: "3.11"

      - name: Install dependencies
        run: pip install python-dotenv

      - name: Restore sent_log from Gist
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          curl -sf -H "Authorization: token $GIST_TOKEN" \\
            "https://api.github.com/gists/$GIST_ID" \\
            | python3 -c "
          import sys, json
          data = json.load(sys.stdin)
          content = data['files'].get('{slug}_sent_log.csv', {{}}).get('content', '')
          print(content, end='')
          " > saas-dev/projects/{slug}/outreach/sent_log.csv || echo "Gist取得失敗"

      - name: Send follow-up emails
        env:
          GMAIL_ADDRESS: ${{{{ secrets.GMAIL_ADDRESS }}}}
          GMAIL_APP_PASSWORD: ${{{{ secrets.GMAIL_APP_PASSWORD }}}}
          SENDER_ADDRESS: ${{{{ secrets.SENDER_ADDRESS }}}}
        run: |
          ARGS="--project {slug} --days ${{{{ inputs.days || '7' }}}}"
          [ "${{{{ inputs.dry_run }}}}" = "true" ] && ARGS="$ARGS --dry-run"
          python shared/gtm/outreach/follow_up.py $ARGS

      - name: Upload updated sent_log to Gist
        if: always()
        env:
          GIST_TOKEN: ${{{{ secrets.GIST_TOKEN }}}}
          GIST_ID: ${{{{ secrets.{slug_upper}_SENT_LOG_GIST_ID }}}}
        run: |
          [ -f saas-dev/projects/{slug}/outreach/sent_log.csv ] || exit 0
          python3 -c "
          import json
          content = open('saas-dev/projects/{slug}/outreach/sent_log.csv', encoding='utf-8').read()
          print(json.dumps({{'files': {{'{slug}_sent_log.csv': {{'content': content}}}}}}))
          " | curl -sf -X PATCH \\
            -H "Authorization: token $GIST_TOKEN" \\
            -H "Content-Type: application/json" \\
            -d @- \\
            "https://api.github.com/gists/$GIST_ID"
'''

    # feedback-pdca
    feedback_pdca = f'''name: {slug}-feedback-pdca

on:
  schedule:
    - cron: "0 20 * * 0"  # 毎週日曜 UTC 20:00（月曜 JST 5:00）
  workflow_dispatch:

jobs:
  pdca:
    runs-on: ubuntu-latest
    permissions:
      contents: write

    steps:
      - uses: actions/checkout@v4

      - uses: actions/setup-python@v5
        with:
          python-version: "3.11"

      - name: Install dependencies
        run: pip install python-dotenv

      - name: Fetch and analyze feedback
        env:
          GEMINI_API_KEY: ${{{{ secrets.GEMINI_API_KEY }}}}
          SUPABASE_URL: ${{{{ secrets.{slug_upper}_SUPABASE_URL }}}}
          SUPABASE_ANON_KEY: ${{{{ secrets.{slug_upper}_SUPABASE_ANON_KEY }}}}
        run: python shared/gtm/outreach/analyze_feedback.py --project {slug}

      - name: Commit report
        run: |
          git config user.email "actions@github.com"
          git config user.name "Auto Bot"
          git add saas-dev/projects/{slug}/feedback_reports/ || true
          git diff --staged --quiet || git commit -m "{slug}: フィードバック分析レポート $(date -u '+%Y-%m-%d') [skip ci]"
          git pull --rebase -X theirs
          git push

      - name: Apply PDCA improvements to gen.py
        env:
          GEMINI_API_KEY: ${{{{ secrets.GEMINI_API_KEY }}}}
          FUDOTEXT_PAT: ${{{{ secrets.FUDOTEXT_PAT }}}}
        run: python shared/gtm/scripts/apply_pdca.py --project {slug}
'''

    # send-code
    send_code = f'''name: {slug}-send-code

on:
  workflow_dispatch:
    inputs:
      to_email:
        description: 'Customer email'
        required: true
      plan:
        description: 'Plan (standard or pro)'
        required: true
      code:
        description: 'Access code (UUID)'
        required: true

jobs:
  send-code:
    runs-on: ubuntu-latest
    steps:
      - name: Send access code email
        env:
          GMAIL_ADDRESS: ${{{{ secrets.GMAIL_ADDRESS }}}}
          GMAIL_APP_PASSWORD: ${{{{ secrets.GMAIL_APP_PASSWORD }}}}
          TO_EMAIL: ${{{{ inputs.to_email }}}}
          PLAN: ${{{{ inputs.plan }}}}
          CODE: ${{{{ inputs.code }}}}
        run: |
          python3 << 'PYEOF'
          import smtplib, os
          from email.mime.text import MIMEText

          to_email = os.environ['TO_EMAIL']
          plan = os.environ['PLAN']
          code = os.environ['CODE']
          gmail = os.environ['GMAIL_ADDRESS']
          pwd = os.environ['GMAIL_APP_PASSWORD']
          plan_label = 'スタンダード' if plan == 'standard' else 'プロ'
          limit = '50件' if plan == 'standard' else '200件'
          price = '¥8,980' if plan == 'standard' else '¥19,800'

          body = (
              f'この度は{product_name} {{plan_label}}プランにお申し込みいただき、誠にありがとうございます。\\n\\n'
              f'アクセスコードをお送りします。\\n\\n'
              f'■ アクセスコード\\n{{code}}\\n\\n'
              f'■ ご利用開始の手順\\n'
              f'1. {product_name}を開く\\n'
              f'   https://{slug}.streamlit.app\\n\\n'
              f'2. メールアドレスを入力してログイン\\n\\n'
              f'3. 無料枠（5件）を使い切ると「アクセスコードを持っている方」欄が表示されます\\n\\n'
              f'4. 上記コードを入力 → {{plan_label}}プラン（月{{limit}}）が有効化されます\\n\\n'
              f'■ プラン内容\\n'
              f'プラン: {{plan_label}}\\n'
              f'月額: {{price}}\\n'
              f'月間上限: {{limit}}\\n\\n'
              f'ご不明な点はお気軽にご返信ください。\\n\\n'
              f'━━━━━━━━━━━━━━━━━━\\n'
              f'TextSeries\\n'
              f'Mail: ryuumg03@gmail.com\\n'
              f'━━━━━━━━━━━━━━━━━━'
          )

          msg = MIMEText(body, 'plain', 'utf-8')
          msg['Subject'] = f'【{product_name}】{{plan_label}}プランのアクセスコードをお送りします'
          msg['From'] = gmail
          msg['To'] = to_email

          with smtplib.SMTP_SSL('smtp.gmail.com', 465, timeout=30) as smtp:
              smtp.login(gmail, pwd)
              smtp.send_message(msg)
              print(f'送信成功: {{to_email}} / {{plan_label}}')
          PYEOF
'''

    wf_dir = dest / ".github" / "workflows"
    _write(wf_dir / f"{slug}-daily-send.yml", daily_send, dry_run)
    _write(wf_dir / f"{slug}-check-replies.yml", check_replies, dry_run)
    _write(wf_dir / f"{slug}-follow-up.yml", follow_up, dry_run)
    _write(wf_dir / f"{slug}-feedback-pdca.yml", feedback_pdca, dry_run)
    _write(wf_dir / f"{slug}-send-code.yml", send_code, dry_run)


def write_email_templates(slug: str, product_name: str,
                          spec: dict, dest: Path, dry_run: bool) -> None:
    email_spec = spec.get("email", {})
    seq1 = email_spec.get("sequence_1",
        f"突然のご連絡をご容赦ください。\n\n"
        f"書類作成業務でお時間を取られていませんか？\n\n"
        f"{product_name}は、情報を入力するだけでAIが30秒で書類のドラフトを生成するサービスです。\n\n"
        f"まずは5件まで無料でお試しいただけます。\n\n"
        f"▶ {product_name}を試してみる\n"
        f"https://{slug}.streamlit.app\n\n"
        f"ご不明な点はお気軽にご返信ください。\n\n"
        f"━━━━━━━━━━━━━━━━━━\n"
        f"TextSeries\n"
        f"Mail: ryuumg03@gmail.com\n"
        f"━━━━━━━━━━━━━━━━━━"
    )
    seq2 = email_spec.get("sequence_2",
        f"先日ご連絡しました{product_name}の件、いかがでしょうか。\n\n"
        f"実際にお試しいただいたユーザーからは「書類作成時間が大幅に短縮された」「品質が安定した」\n"
        f"というお声をいただいております。\n\n"
        f"引き続き5件まで無料でご利用いただけます。\n\n"
        f"▶ {product_name}を試してみる\n"
        f"https://{slug}.streamlit.app\n\n"
        f"━━━━━━━━━━━━━━━━━━\n"
        f"TextSeries\n"
        f"Mail: ryuumg03@gmail.com\n"
        f"━━━━━━━━━━━━━━━━━━"
    )
    tpl_dir = dest / "shared" / "gtm" / "outreach" / "templates"
    _write(tpl_dir / f"{slug}_sequence_1.txt", seq1, dry_run)
    _write(tpl_dir / f"{slug}_sequence_2.txt", seq2, dry_run)


def write_outreach_scripts(slug: str, industry: str, spec: dict,
                           dest: Path, dry_run: bool) -> None:
    """fetch_leads.py / pipeline.py / send_emails.py を生成する。"""
    outreach_spec = spec.get("outreach", {})
    icp = spec.get("icp", {})

    queries_raw = outreach_spec.get("search_queries", [])
    company_kws = outreach_spec.get("company_required_keywords",
                                    icp.get("company_name_hint_keywords", []))
    blog_signals = outreach_spec.get("blog_signals",
                                     ["コツ", "選び方", "ランキング", "比較", "まとめ", "一覧"])
    # Clean up placeholder text Gemini sometimes puts in lists
    queries = [q for q in queries_raw
               if q and not q.startswith("（") and len(q) > 10]
    if not queries:
        kws = icp.get("target_keywords", [industry])
        queries = [
            f'{kws[0] if kws else industry} "お問い合わせ" site:co.jp -求人',
            f'{kws[0] if kws else industry} "info@" -採用',
            f'{kws[0] if kws else industry} "メールでのお問い合わせ" -転職',
        ]

    queries_repr = repr(queries)
    company_kws_repr = repr(company_kws)
    blog_signals_repr = repr(blog_signals)

    # ── fetch_leads.py ──────────────────────────────────────────────────────
    fetch_leads = f'''"""
{spec.get("product_name", slug)} GTMリード収集スクリプト（Brave API）
  python fetch_leads.py [--limit 150]
"""
import csv
import json
import os
import re
import sys
import time
import urllib.parse
import urllib.request
from datetime import datetime
from pathlib import Path

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, Exception):
        pass

_DIR = Path(__file__).parent
LEADS_FILE = _DIR / "leads.csv"

try:
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=_DIR.parent.parent.parent.parent / ".env")
except ImportError:
    pass

BRAVE_KEY = os.environ.get("BRAVE_API_KEY", "")
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\\-]+@[a-zA-Z0-9.\\-]+\\.[a-zA-Z]{{2,}}")
OG_SITE_RE = re.compile(
    r\'<meta[^>]+property=["\\\'"]og:site_name["\\\'"][^>]+content=["\\\'"]([^"\\\' ]{{2,40}})["\\\'"]\'
    r\'|<meta[^>]+content=["\\\'"]([^"\\\' ]{{2,40}})["\\\'"][^>]+property=["\\\'"]og:site_name["\\\'"]\',
    re.IGNORECASE,
)
TITLE_RE = re.compile(r"<title[^>]*>([^<]+)</title>", re.IGNORECASE)

EMAIL_SKIP = ["noreply", "no-reply", "example", "sentry", "google",
              "schema.org", "w3.org", "placeholder", "sample@", "test@",
              "@sample.", "@mail.jp", "@example.", "postmaster@", "webmaster@"]
FAKE_TLDS = {{".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".ico", ".pdf", ".zip"}}
SITE_SKIP = ["wikipedia", "google", "yahoo", "twitter", "facebook", "instagram",
             "amazon", "rakuten", "mynavi", "doda", "rikunabi", "indeed",
             "townwork", "hellowork", "nikkei", "nhk", "pref.", "city.", "go.jp"]

_COMPANY_KEYWORDS = {company_kws_repr}

QUERIES = {queries_repr}


def _brave_search(query: str, count: int = 10) -> list[dict]:
    if not BRAVE_KEY:
        return []
    url = f"https://api.search.brave.com/res/v1/web/search?q={{urllib.parse.quote(query)}}&count={{count}}&country=jp"
    try:
        req = urllib.request.Request(url, headers={{
            "Accept": "application/json",
            "X-Subscription-Token": BRAVE_KEY,
        }})
        with urllib.request.urlopen(req, timeout=15) as r:
            data = json.loads(r.read())
        return data.get("web", {{}}).get("results", [])
    except Exception as e:
        print(f"  Brave API エラー: {{e}}")
        return []


def _fetch_html(url: str) -> str:
    try:
        req = urllib.request.Request(url, headers={{"User-Agent": "Mozilla/5.0 (compatible; TextSeriesBot/1.0)"}})
        with urllib.request.urlopen(req, timeout=10) as r:
            raw = r.read()
            enc = r.headers.get_content_charset("utf-8")
            return raw.decode(enc, errors="ignore")
    except Exception:
        return ""


def _extract_emails(html: str) -> list[str]:
    found = EMAIL_RE.findall(html)
    result = []
    for e in found:
        e = e.lower()
        if any(b in e for b in EMAIL_SKIP):
            continue
        if any(e.endswith(t) for t in FAKE_TLDS):
            continue
        if e not in result:
            result.append(e)
    return result[:2]


def _extract_company_name(html: str, fallback: str = "") -> str:
    m = OG_SITE_RE.search(html)
    if m:
        name = (m.group(1) or m.group(2) or "").strip()
        if name and any(kw in name for kw in _COMPANY_KEYWORDS):
            return name[:40]
    t = TITLE_RE.search(html)
    if t:
        title = t.group(1).strip()
        for sep in ["｜", "|", "–", "-", "—", "　"]:
            parts = title.split(sep)
            if len(parts) == 1:
                continue
            for part in parts:
                part = part.strip()
                if any(kw in part for kw in _COMPANY_KEYWORDS) and len(part) <= 30:
                    return part
    return fallback[:40] if fallback else ""


def main(limit: int = 150):
    if not BRAVE_KEY:
        print("BRAVE_API_KEY が未設定です。.envに追加してください。")
        return

    existing = set()
    if LEADS_FILE.exists():
        with open(LEADS_FILE, encoding="utf-8") as f:
            existing = {{row["url"] for row in csv.DictReader(f)}}
    print(f"既存リード: {{len(existing)}}件")

    write_header = not LEADS_FILE.exists()
    collected = 0

    with open(LEADS_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["company_name", "email", "url", "prefecture", "scraped_at"])
        if write_header:
            writer.writeheader()

        for query in QUERIES:
            if collected >= limit:
                break
            print(f"\\nクエリ: {{query[:60]}}...")
            results = _brave_search(query, count=10)
            time.sleep(1.0)

            for r in results:
                if collected >= limit:
                    break
                url = r.get("url", "")
                if not url or any(s in url for s in SITE_SKIP):
                    continue
                if url in existing:
                    continue

                html = _fetch_html(url)
                if not html:
                    time.sleep(0.5)
                    continue

                emails = _extract_emails(html)
                if not emails:
                    for path in ["/contact", "/inquiry", "/contact.html"]:
                        from urllib.parse import urlparse
                        base = f"{{urlparse(url).scheme}}://{{urlparse(url).netloc}}"
                        contact_html = _fetch_html(base + path)
                        if contact_html:
                            emails = _extract_emails(contact_html)
                            if emails:
                                break
                        time.sleep(0.3)

                if not emails:
                    existing.add(url)
                    continue

                company = _extract_company_name(html, r.get("title", ""))
                if not company:
                    existing.add(url)
                    continue

                for email in emails:
                    writer.writerow({{
                        "company_name": company,
                        "email": email,
                        "url": url,
                        "prefecture": "",
                        "scraped_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    }})
                    f.flush()
                    print(f"  + {{company[:30]}} | {{email}}")

                existing.add(url)
                collected += 1
                time.sleep(1.0)

    print(f"\\n完了。{{collected}}件収集しました -> {{LEADS_FILE}}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=150)
    args = parser.parse_args()
    main(args.limit)
'''

    # ── pipeline.py ─────────────────────────────────────────────────────────
    pipeline = f'''"""
{spec.get("product_name", slug)} GTMパイプライン（fetch → qualify → generate → ready to send）
  python pipeline.py [--limit 50]
"""
import argparse
import subprocess
import sys
from pathlib import Path

PROJECT = "{slug}"
_DIR = Path(__file__).parent
_ROOT = _DIR.parent.parent.parent


def run(cmd: list[str]) -> int:
    print(f"\\n{{"="*60}}")
    print(f"実行: {{' '.join(cmd)}}")
    print("="*60)
    result = subprocess.run(cmd, cwd=str(_ROOT))
    return result.returncode


def main(limit: int = 50):
    steps = [
        [sys.executable, str(_DIR / "fetch_leads.py"), "--limit", str(limit)],
        [sys.executable, str(_ROOT / "shared/gtm/leads/qualify_leads.py"),
         "--project", PROJECT, "--input", str(_DIR / "leads.csv")],
        [sys.executable, str(_ROOT / "shared/gtm/outreach/generate_emails.py"),
         "--project", PROJECT, "--limit", str(limit)],
    ]
    for cmd in steps:
        rc = run(cmd)
        if rc != 0:
            print(f"\\nエラー: {{cmd[1]}} が終了コード {{rc}} で失敗しました。")
            sys.exit(rc)
    print(f"\\n✅ パイプライン完了。send_emails.py --dry-run で内容を確認してください。")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=50)
    args = parser.parse_args()
    main(args.limit)
'''

    # ── send_emails.py ───────────────────────────────────────────────────────
    send_emails = f'''"""
{spec.get("product_name", slug)} アウトリーチ送信スクリプト（Gmail SMTP・30件/日）
"""
import argparse
import csv
import os
import smtplib
import sys
import time
from datetime import datetime, timezone
from email.mime.text import MIMEText
from pathlib import Path

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, Exception):
        pass

_DIR = Path(__file__).parent
DRAFT_FILE = _DIR / "emails_draft.csv"
SENT_LOG = _DIR / "sent_log.csv"
OPT_OUT_FILE = _DIR / "opt_out.csv"

try:
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=_DIR.parent.parent.parent.parent / ".env")
except ImportError:
    pass

GMAIL_ADDRESS = os.environ.get("GMAIL_ADDRESS", "")
GMAIL_APP_PASSWORD = os.environ.get("GMAIL_APP_PASSWORD", "")

DAILY_LIMIT = 30
SEND_INTERVAL = 30

_COMPANY_REQUIRED = {company_kws_repr}
_BLOG_SIGNALS = {blog_signals_repr}


def _send(to: str, subject: str, body: str) -> bool:
    try:
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = GMAIL_ADDRESS
        msg["To"] = to
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=30) as smtp:
            smtp.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
            smtp.send_message(msg)
        return True
    except Exception as e:
        print(f"    送信失敗: {{e}}")
        return False


def _load_sent() -> set[str]:
    if not SENT_LOG.exists():
        return set()
    with open(SENT_LOG, encoding="utf-8", newline="") as f:
        return {{row["email"] for row in csv.DictReader(f)}}


def _load_opt_out() -> set[str]:
    if not OPT_OUT_FILE.exists():
        return set()
    with open(OPT_OUT_FILE, encoding="utf-8", newline="") as f:
        return {{row["email"] for row in csv.DictReader(f)}}


def add_opt_out(email: str, reason: str = "") -> None:
    existing = _load_opt_out()
    if email in existing:
        return
    write_header = not OPT_OUT_FILE.exists()
    with open(OPT_OUT_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["email", "reason", "added_at"])
        if write_header:
            writer.writeheader()
        writer.writerow({{"email": email, "reason": reason, "added_at": datetime.now().strftime("%Y-%m-%d %H:%M")}})
    print(f"opt_out登録完了: {{email}}")


def _today_sent_count() -> int:
    if not SENT_LOG.exists():
        return 0
    today = datetime.now().strftime("%Y-%m-%d")
    with open(SENT_LOG, encoding="utf-8", newline="") as f:
        return sum(1 for row in csv.DictReader(f) if row.get("sent_at", "").startswith(today))


def _check_safety(drafts: list[dict]) -> bool:
    ok = True
    for d in drafts:
        name = d.get("company_name", "")
        if any(sig in name for sig in _BLOG_SIGNALS):
            print(f"WARNING: {{d['email']}} — 会社名がブログタイトルの可能性: 「{{name[:30]}}」")
            ok = False
        elif name and _COMPANY_REQUIRED and not any(kw in name for kw in _COMPANY_REQUIRED):
            print(f"WARNING: {{d['email']}} — 対象業種キーワードなし: 「{{name[:30]}}」")
            ok = False
        body = d.get("body", "")
        missing = []
        if "TextSeries" not in body:
            missing.append("送信者名")
        if not any(p in body for p in ["配信停止", "送付を停止", "送信停止", "ご連絡いたしません"]):
            missing.append("配信停止文言")
        if "〒" not in body and "住所" not in body:
            missing.append("住所")
        if missing:
            print(f"WARNING: {{d['email']}} — 特定電子メール法違反の可能性: {{', '.join(missing)}}")
            ok = False
    return ok


def main(limit: int = DAILY_LIMIT, dry_run: bool = False, test_to: str = "", force_send: bool = False):
    if not dry_run and not test_to:
        now_jst = (datetime.now(timezone.utc).hour + 9) % 24
        if not force_send and not (9 <= now_jst < 18):
            print(f"送信停止: 現在 JST {{now_jst}}時台です（送信は JST 9:00〜18:00）。--force-send で強制送信。")
            return

    if not DRAFT_FILE.exists():
        print("emails_draft.csv が見つかりません。generate_emails.py を先に実行してください。")
        return

    with open(DRAFT_FILE, encoding="utf-8", newline="") as f:
        drafts = list(csv.DictReader(f))

    if test_to:
        if not GMAIL_ADDRESS or not GMAIL_APP_PASSWORD:
            print("環境変数 GMAIL_ADDRESS / GMAIL_APP_PASSWORD が未設定です。")
            return
        draft = next((d for d in drafts if d.get("status") == "draft"), None) or (drafts[0] if drafts else None)
        if not draft:
            print("ドラフトが1件もありません。")
            return
        print(f"テスト送信: {{test_to}}")
        ok = _send(test_to, f"【テスト】{{draft['subject']}}", draft["body"])
        print("成功" if ok else "失敗")
        return

    if dry_run:
        already_sent = _load_sent()
        targets = [d for d in drafts if d.get("status") == "draft" and d["email"] not in already_sent][:limit]
        print(f"\\nDRY-RUN: 送信予定 {{len(targets)}}件")
        for i, d in enumerate(targets, 1):
            print(f"[{{i:2}}] {{d['company_name'][:30]:<30}} | {{d['email']}}")
            print(f"     件名: {{d['subject']}}")
        return

    if not GMAIL_ADDRESS or not GMAIL_APP_PASSWORD:
        print("環境変数が未設定: GMAIL_ADDRESS / GMAIL_APP_PASSWORD")
        return

    already_sent = _load_sent()
    opt_out = _load_opt_out()
    today_count = _today_sent_count()
    remaining = limit - today_count

    if remaining <= 0:
        print(f"本日の送信上限（{{limit}}件）に達しています。")
        return

    print(f"本日送信済み: {{today_count}}件 / 残り: {{remaining}}件")
    targets = [d for d in drafts if d.get("status") == "draft"
               and d["email"] not in already_sent and d["email"] not in opt_out]
    targets = targets[:remaining]
    print(f"送信対象: {{len(targets)}}件")

    if not targets:
        print("送信対象がありません。pipeline.py を実行してリードを補充してください。")
        return

    if not _check_safety(targets):
        print("\\n安全チェックに失敗しました。--dry-run で内容を確認してください。")
        return

    write_header = not SENT_LOG.exists()
    sent_emails = set()

    with open(SENT_LOG, "a", newline="", encoding="utf-8") as log_f:
        log_writer = csv.DictWriter(log_f, fieldnames=["company_name", "email", "subject", "sent_at", "result"])
        if write_header:
            log_writer.writeheader()

        for i, draft in enumerate(targets, 1):
            email = draft["email"]
            print(f"  [{{i}}/{{len(targets)}}] {{draft['company_name'][:30]}} <{{email}}>", end=" ... ")
            ok = _send(email, draft["subject"], draft["body"])
            result = "sent" if ok else "failed"
            print(result)
            log_writer.writerow({{
                "company_name": draft["company_name"],
                "email": email,
                "subject": draft["subject"],
                "sent_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "result": result,
            }})
            log_f.flush()
            if ok:
                sent_emails.add(email)
            if i < len(targets):
                time.sleep(SEND_INTERVAL)

    draft_fields = list(drafts[0].keys()) if drafts else ["company_name", "email", "subject", "body", "url", "status"]
    updated = [{{**row, "status": "sent"}} if row["email"] in sent_emails else row for row in drafts]

    with open(DRAFT_FILE, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=draft_fields)
        writer.writeheader()
        writer.writerows(updated)

    print(f"\\n完了。送信成功: {{len(sent_emails)}}件")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=DAILY_LIMIT)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--test-to", type=str, default="")
    parser.add_argument("--force-send", action="store_true")
    parser.add_argument("--add-opt-out", type=str, default="")
    parser.add_argument("--opt-out-reason", type=str, default="返信による申し出")
    args = parser.parse_args()

    if args.add_opt_out:
        add_opt_out(args.add_opt_out, args.opt_out_reason)
    else:
        main(limit=args.limit, dry_run=args.dry_run, test_to=args.test_to, force_send=args.force_send)
'''

    outreach_dir = dest / "saas-dev" / "projects" / slug / "outreach"
    feedback_dir = dest / "saas-dev" / "projects" / slug / "feedback_reports"
    _write(outreach_dir / "fetch_leads.py", fetch_leads, dry_run)
    _write(outreach_dir / "pipeline.py", pipeline, dry_run)
    _write(outreach_dir / "send_emails.py", send_emails, dry_run)
    if not dry_run:
        feedback_dir.mkdir(parents=True, exist_ok=True)
        gitkeep = feedback_dir / ".gitkeep"
        if not gitkeep.exists():
            gitkeep.write_text("")
    print(f"    ✓ dir feedback_reports/")


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="TextSeries 新製品ジェネレーター")
    parser.add_argument("--slug", required=True, help="製品スラッグ (例: rifotext)")
    parser.add_argument("--industry", required=True, help="ターゲット業界 (例: リフォーム会社)")
    parser.add_argument("--docs", required=True, help="書類種別カンマ区切り (例: 工事見積書,提案書)")
    parser.add_argument("--emoji", required=True, help="絵文字 (例: 🔨)")
    parser.add_argument("--rank", type=int, default=99, help="製品ランク")
    parser.add_argument("--product-name", help="製品名 (省略時: slug をCapitalize)")
    parser.add_argument("--no-stripe", action="store_true", help="setup_stripe.py をスキップ")
    parser.add_argument("--dry-run", action="store_true", help="ファイル書き込みをスキップ")
    args = parser.parse_args()

    slug = args.slug.lower().strip()
    industry = args.industry.strip()
    docs = [d.strip() for d in args.docs.split(",") if d.strip()]
    emoji = args.emoji.strip()
    product_name = args.product_name or (slug[0].upper() + slug[1:])

    if not docs:
        sys.exit("--docs に書類種別を1つ以上指定してください")

    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key:
        try:
            from dotenv import load_dotenv
            load_dotenv(dotenv_path=Path("C:/Users/ryuuM/fudotext/.env"))
            api_key = os.environ.get("GEMINI_API_KEY", "")
        except ImportError:
            pass
    if not api_key:
        sys.exit("GEMINI_API_KEY が環境変数または .env に設定されていません")

    print(f"\n🚀 {product_name} ({slug}) の生成を開始します")
    print(f"   業界: {industry}")
    print(f"   書類: {', '.join(docs)}")
    print(f"   絵文字: {emoji}  ランク: {args.rank}")
    if args.dry_run:
        print("   [dry-run モード: ファイル書き込みなし]")

    # 1. Generate spec via Gemini
    spec = generate_spec(slug, industry, docs, emoji, product_name, api_key)
    print(f"  ✓ 仕様生成完了: {spec.get('tagline', '')}")

    # 2. Write fudotext files
    fudotext_dir = FUDOTEXT / "saas-dev" / "projects" / slug
    print("\n📁 fudotext アプリファイルを生成中...")
    write_requirements_txt(fudotext_dir, args.dry_run)
    write_db_py(slug, fudotext_dir, args.dry_run)
    write_gen_py(slug, product_name, spec, docs, fudotext_dir, args.dry_run)
    write_app_py(slug, product_name, emoji, spec, docs, fudotext_dir, args.dry_run)

    # 3. Write ai-holdings files
    print("\n📁 GTM・ワークフローファイルを生成中...")
    write_config_json(slug, product_name, args.rank, spec, docs, AI_HOLDINGS, args.dry_run)
    write_lp_html(slug, product_name, emoji, spec, docs, AI_HOLDINGS, args.dry_run)
    write_workflows(slug, product_name, AI_HOLDINGS, args.dry_run)
    write_email_templates(slug, product_name, spec, AI_HOLDINGS, args.dry_run)
    write_outreach_scripts(slug, industry, spec, AI_HOLDINGS, args.dry_run)

    # 4. Run setup_stripe.py
    if not args.no_stripe and not args.dry_run:
        print("\n⚙️  setup_stripe.py を実行中...")
        setup_script = AI_HOLDINGS / "shared" / "gtm" / "scripts" / "setup_stripe.py"
        result = subprocess.run(
            [sys.executable, str(setup_script), "--project", slug],
            capture_output=False,
        )
        if result.returncode != 0:
            print(f"  ⚠️  setup_stripe.py がエラーで終了しました (code={result.returncode})")
            print("  clipboard.txt・Stripe・GCP設定を手動で確認してください")
    elif args.no_stripe:
        print("\n⏭️  --no-stripe: setup_stripe.py をスキップ")
        print(f"  手動実行: python shared/gtm/scripts/setup_stripe.py --project {slug}")

    print(f"\n✅ {product_name} の生成が完了しました！")
    print(f"\n次のステップ:")
    print(f"  1. clipboard.txt の {slug.upper()} セクションを Streamlit Cloud secrets に貼り付け")
    print(f"  2. {FUDOTEXT / 'saas-dev' / 'projects' / slug} を Streamlit Cloud にデプロイ")
    print(f"  3. saas-dev/projects/{slug}/outreach/ に pipeline.py・send_emails.py を配置")
    print(f"  4. GitHub Secrets に {slug.upper()}_SENT_LOG_GIST_ID を追加")
    print(f"  5. LP: https://ryuu321.github.io/ai-holdings/docs/{slug}.html")


if __name__ == "__main__":
    main()
